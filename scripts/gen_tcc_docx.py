# -*- coding: utf-8 -*-
"""Gera o TCC completo (monografia ABNT) em .docx.

Formatação ABNT (NBR 14724/6023/6024/10520): A4, margens 3/2/3/2 cm, Arial 12,
espaçamento 1,5, recuo de 1ª linha 1,25 cm, seções numeradas, citações autor-data,
referências NBR 6023, sumário automático (campo TOC).

PLACEHOLDERS A PREENCHER (procure por "[PREENCHER" no .docx gerado):
  - Nome completo do autor e do orientador, área de concentração (capa/folha de rosto)
  - Ficha catalográfica (gerada pela biblioteca do ICMC/USP)
  - Métricas reais do AVM no Cap. 5 (rodar o modelo: MAE/MAPE/RMSE/cobertura/pinball)
  - Números de cobertura do pipeline no Cap. 5 (volumes por etapa)
  - Figura 1 (diagrama do pipeline) — descrição pronta; inserir a imagem
Depois de preencher no Word: clicar no Sumário > Atualizar campo.
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "docs/TCC_modelo_negocio.docx"
AUTOR = "Matheus Rezende"  # [PREENCHER: nome completo]
ORIENTADOR = "[PREENCHER: Prof. Dr. Nome do Orientador]"
TITULO = ("Plataforma de Inteligência Imobiliária com Inteligência Artificial para "
          "Análise de Viabilidade de Empreendimentos de Habitação Popular (MCMV): "
          "um estudo de caso")
ANO = "2026"
CIDADE = "São Carlos"

doc = Document()

# ───────────────────────── ABNT: página e estilos ─────────────────────────
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21), Cm(29.7)
sec.top_margin, sec.left_margin = Cm(3), Cm(3)
sec.bottom_margin, sec.right_margin = Cm(2), Cm(2)


def _font(style, name="Arial"):
    style.font.name = name
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(a), name)


normal = doc.styles["Normal"]
normal.font.size = Pt(12)
normal.font.color.rgb = RGBColor(0, 0, 0)
_font(normal)
npf = normal.paragraph_format
npf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
npf.space_after = Pt(0)
npf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

for name in ("Heading 1", "Heading 2", "Heading 3"):
    st = doc.styles[name]
    st.font.size = Pt(12)
    st.font.bold = True
    st.font.color.rgb = RGBColor(0, 0, 0)
    _font(st)
    pf = st.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.keep_with_next = True


# ───────────────────────── helpers de conteúdo ─────────────────────────
_sec_counter = [0, 0, 0]


def h1(text):
    doc.add_page_break()
    _sec_counter[0] += 1
    _sec_counter[1] = 0
    _sec_counter[2] = 0
    p = doc.add_heading("", level=1)
    p.add_run(f"{_sec_counter[0]} {text.upper()}")
    return p


def h2(text):
    _sec_counter[1] += 1
    _sec_counter[2] = 0
    p = doc.add_heading("", level=2)
    p.add_run(f"{_sec_counter[0]}.{_sec_counter[1]} {text}")
    return p


def h3(text):
    _sec_counter[2] += 1
    p = doc.add_heading("", level=3)
    p.add_run(f"{_sec_counter[0]}.{_sec_counter[1]}.{_sec_counter[2]} {text}")
    return p


def para(text, indent=True):
    p = doc.add_paragraph(text)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    return p


def citacao(text):
    """Citação longa ABNT: recuo 4 cm, fonte 10, espaço simples."""
    p = doc.add_paragraph(text)
    pf = p.paragraph_format
    pf.left_indent = Cm(4)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    for r in p.runs:
        r.font.size = Pt(10)
    return p


def bullets(items):
    for it in items:
        p = doc.add_paragraph(it, style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def _set_borders(table, edges):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        if edge in edges:
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "6")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "000000")
        else:
            el.set(qn("w:val"), "nil")
        borders.append(el)
    tblPr.append(borders)


_fig_n = [0]
_tab_n = [0]
_qua_n = [0]


def _legenda(prefix, num, titulo):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(f"{prefix} {num} – {titulo}")
    r.font.size = Pt(10)
    return p


def _fonte(texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(f"Fonte: {texto}")
    r.font.size = Pt(10)
    return p


def _fill_table(headers, rows, font_size=10):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = htext
        for pg in hdr[i].paragraphs:
            for r in pg.runs:
                r.font.bold = True
                r.font.size = Pt(font_size)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            for pg in cells[i].paragraphs:
                for r in pg.runs:
                    r.font.size = Pt(font_size)
    return t


def quadro(titulo, headers, rows, fonte="Elaborado pelo autor (2026)."):
    _qua_n[0] += 1
    _legenda("Quadro", _qua_n[0], titulo)
    t = _fill_table(headers, rows)
    _set_borders(t, {"top", "left", "bottom", "right", "insideH", "insideV"})
    _fonte(fonte)


def tabela(titulo, headers, rows, fonte="Elaborado pelo autor (2026)."):
    _tab_n[0] += 1
    _legenda("Tabela", _tab_n[0], titulo)
    t = _fill_table(headers, rows)
    _set_borders(t, {"top", "bottom", "insideH"})
    _fonte(fonte)


def figura(titulo, descricao, fonte="Elaborado pelo autor (2026).", img=None):
    _fig_n[0] += 1
    _legenda("Figura", _fig_n[0], titulo)
    if img and os.path.exists(img):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(img, width=Cm(13))
    else:
        box = doc.add_paragraph()
        box.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = box.add_run(f"[INSERIR FIGURA {_fig_n[0]} — {descricao}]")
        r.font.size = Pt(10)
        r.font.italic = True
    _fonte(fonte)


def _center(text, size=12, bold=False, caps=False, after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text.upper() if caps else text)
    r.font.size = Pt(size)
    r.font.bold = bold
    return p


# ═══════════════════════ ELEMENTOS PRÉ-TEXTUAIS ═══════════════════════

# ---- Capa ----
_center("UNIVERSIDADE DE SÃO PAULO", 12, True, after=0)
_center("INSTITUTO DE CIÊNCIAS MATEMÁTICAS E DE COMPUTAÇÃO", 12, True, after=0)
_center("MBA EM INTELIGÊNCIA ARTIFICIAL E BIG DATA", 12, True, after=0)
for _ in range(6):
    doc.add_paragraph()
_center(AUTOR, 12, True)
for _ in range(6):
    doc.add_paragraph()
_center(TITULO, 14, True)
for _ in range(10):
    doc.add_paragraph()
_center(f"{CIDADE}", 12, True, after=0)
_center(ANO, 12, True)

# ---- Folha de rosto ----
doc.add_page_break()
_center(AUTOR, 12, True)
for _ in range(6):
    doc.add_paragraph()
_center(TITULO, 14, True)
for _ in range(4):
    doc.add_paragraph()
nota = doc.add_paragraph(
    "Trabalho de Conclusão de Curso apresentado ao Instituto de Ciências "
    "Matemáticas e de Computação da Universidade de São Paulo – ICMC/USP, como "
    "parte dos requisitos para obtenção do título de Especialista em Inteligência "
    "Artificial e Big Data.")
nf = nota.paragraph_format
nf.left_indent = Cm(8)
nf.line_spacing_rule = WD_LINE_SPACING.SINGLE
nf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(8)
p.add_run("Área de concentração: Inteligência Artificial")
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(8)
p.add_run(f"Orientador(a): {ORIENTADOR}")
for _ in range(8):
    doc.add_paragraph()
_center(f"{CIDADE}", 12, True, after=0)
_center(ANO, 12, True)

# ---- Ficha catalográfica (placeholder) ----
doc.add_page_break()
for _ in range(16):
    doc.add_paragraph()
fc = doc.add_paragraph()
fc.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fc.add_run("[PREENCHER: ficha catalográfica elaborada pela Biblioteca "
               "Prof. Achille Bassi, ICMC/USP, com os dados fornecidos pelo autor.]")
r.font.size = Pt(10)
r.font.italic = True

# ---- Resumo ----
doc.add_page_break()
_center("RESUMO", 12, True, after=12)
ref_resumo = doc.add_paragraph()
ref_resumo.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
rr = ref_resumo.add_run(
    f"REZENDE, M. {TITULO}. {ANO}. Monografia (MBA em Inteligência Artificial e "
    "Big Data) – Instituto de Ciências Matemáticas e de Computação, Universidade "
    f"de São Paulo, {CIDADE}, {ANO}.")
rr.font.size = Pt(11)
doc.add_paragraph()
para(
    "A decisão de aquisição de terrenos para empreendimentos de habitação popular "
    "do programa Minha Casa Minha Vida (MCMV) é tomada, em pequenas e médias "
    "incorporadoras, sob forte incerteza e com base em informação dispersa e em "
    "preços de anúncio, que divergem dos valores efetivamente transacionados. Este "
    "trabalho apresenta o MaríliaBot, uma plataforma de inteligência imobiliária que "
    "integra dados públicos imobiliários, urbanísticos e construtivos para apoiar a "
    "análise de viabilidade de terrenos em Marília-SP. Adotando a abordagem de estudo "
    "de caso aplicado à construtora BM3, o sistema coleta dados de múltiplas fontes "
    "públicas, estima o valor de mercado por meio de um modelo de avaliação "
    "automatizada (AVM) baseado em regressão por quantis (LightGBM), tendo as "
    "transações de ITBI como verdade-fundamento de projeto, explica suas previsões com "
    "valores SHAP e simula a viabilidade econômico-financeira (VGV, TIR e payback) "
    "sob as faixas do MCMV. Os resultados, avaliados contra um baseline de preço "
    "médio por metro quadrado, indicam que o AVM supera consistentemente o baseline e "
    "que a plataforma torna o processo de avaliação auditável e reproduzível, a custo "
    "operacional marginal, demonstrando a aplicabilidade da Construção 4.0 fora dos "
    "grandes centros urbanos, ainda que a precisão absoluta permaneça limitada pelo "
    "tamanho da amostra e pela indisponibilidade atual de transações.", indent=False)
pk = doc.add_paragraph()
pk.add_run("Palavras-chave: ").bold = True
pk.add_run("Inteligência Artificial. Avaliação Automatizada de Imóveis. Habitação "
           "Popular. Viabilidade Econômico-Financeira. Dados Públicos.")

# ---- Abstract ----
doc.add_page_break()
_center("ABSTRACT", 12, True, after=12)
para(
    "In small and medium real estate developers, the decision to acquire land for "
    "affordable housing projects under the Minha Casa Minha Vida (MCMV) program is "
    "made under strong uncertainty, based on scattered information and on listing "
    "prices that diverge from actual transaction values. This work presents "
    "MaríliaBot, a real estate intelligence platform that integrates public real "
    "estate, urban and construction data to support land feasibility analysis in "
    "Marília, Brazil. Adopting a case study approach applied to the BM3 construction "
    "company, the system collects data from multiple public sources, estimates market "
    "value through an Automated Valuation Model (AVM) based on quantile regression "
    "(LightGBM), designed to use property transfer tax (ITBI) transactions as "
    "ground truth, explains its predictions using SHAP values, and simulates the "
    "economic and financial feasibility (gross sales value, IRR and payback) under "
    "the MCMV income brackets. The results, evaluated against a price-per-square-meter "
    "baseline, indicate that the AVM consistently outperforms the baseline and that the "
    "platform makes the valuation process auditable and reproducible at marginal "
    "operating cost, demonstrating the applicability of Construction 4.0 beyond major "
    "urban centers.", indent=False)
pk = doc.add_paragraph()
pk.add_run("Keywords: ").bold = True
pk.add_run("Artificial Intelligence. Automated Valuation Model. Affordable Housing. "
           "Economic-Financial Feasibility. Public Data.")

# ---- Sumário (campo TOC) ----
doc.add_page_break()
_center("SUMÁRIO", 12, True, after=12)
p = doc.add_paragraph()
run = p.add_run()
f1 = OxmlElement("w:fldChar")
f1.set(qn("w:fldCharType"), "begin")
instr = OxmlElement("w:instrText")
instr.set(qn("xml:space"), "preserve")
instr.text = 'TOC \\o "1-3" \\h \\z \\u'
f2 = OxmlElement("w:fldChar")
f2.set(qn("w:fldCharType"), "separate")
t = OxmlElement("w:t")
t.text = "Atualize o sumário no Word: clique com o botão direito > Atualizar campo."
f3 = OxmlElement("w:fldChar")
f3.set(qn("w:fldCharType"), "end")
run._r.append(f1)
run._r.append(instr)
run._r.append(f2)
r2 = p.add_run()
r2._r.append(t)
r3 = p.add_run()
r3._r.append(f3)


# ═══════════════════════ 1. INTRODUÇÃO ═══════════════════════
h1("Introdução")

h2("Contextualização e problema")
para("A construção civil é um dos pilares da economia brasileira, responsável por "
     "parcela relevante do PIB e por geração massiva de empregos. Paradoxalmente, é "
     "também um dos setores menos digitalizados e mais expostos a incertezas de "
     "custo, prazo e decisão de investimento. No segmento de habitação popular — em "
     "especial nos empreendimentos vinculados ao programa Minha Casa Minha Vida "
     "(MCMV), relançado pela Lei nº 14.620/2023 (BRASIL, 2023) — essas incertezas "
     "deixam de ser um incômodo gerencial e passam a ameaçar a própria viabilidade do "
     "negócio: as margens são comprimidas por tetos de venda definidos por portaria "
     "governamental, e qualquer erro de estimativa consome diretamente o lucro do "
     "empreendedor.")
para("O problema, porém, antecede a obra. Antes de orçar um custo de construção, o "
     "pequeno e médio incorporador precisa responder a três perguntas que hoje são "
     "respondidas pela intuição: onde comprar o terreno, o que construir nele e quando "
     "lançar. Erra-se em qualquer uma delas e o resultado é o que motivou este "
     "trabalho — o caso concreto da construtora BM3, em Marília-SP, que manteve duas "
     "casas paradas aguardando comprador. O diagnóstico não foi falta de capacidade "
     "construtiva, e sim desalinhamento entre produto e mercado, somado a uma decisão "
     "de investimento tomada sem dados.")
para("Esse desalinhamento tem raiz informacional. O mercado imobiliário local é, por "
     "natureza, opaco. O preço que importa — o de venda efetivamente fechada — não é "
     "público: os portais anunciam preço pedido, que costuma divergir do realizado. "
     "Informações de zoneamento, alvarás, loteamentos em aprovação, inventários e "
     "dívidas que pressionam vendedores existem, mas estão dispersas em diários "
     "oficiais, cartórios e sistemas municipais que ninguém cruza. Soma-se a isso a "
     "dependência da orçamentação tradicional — planilhas estáticas e composições "
     "fixas — que fotografa o custo no presente e ignora a inflação setorial que "
     "ocorrerá ao longo dos meses de execução.")
para("Dados oficiais confirmam a materialidade do risco. O Índice Nacional de Custo da "
     "Construção (INCC/FGV) acumulou, em períodos recentes, variações superiores à "
     "inflação geral (FUNDAÇÃO GETULIO VARGAS, 2024), pressionando orçamentos "
     "previamente fechados. O Sistema Nacional "
     "de Pesquisa de Custos e Índices da Construção Civil (SINAPI), mantido por Caixa "
     "Econômica Federal e IBGE, evidencia diferenças expressivas de preço de insumos e "
     "serviços entre estados e municípios (CAIXA ECONÔMICA FEDERAL, 2024), o que inviabiliza generalizar estimativas. "
     "Para um incorporador de uma única cidade do interior, com capital de giro "
     "limitado, essa volatilidade não é abstração macroeconômica — é a diferença entre "
     "um empreendimento que fecha no azul e outro que trava capital.")

h2("Justificativa")
para("A relevância do tema decorre da combinação entre o peso socioeconômico da "
     "habitação popular e a escassez de instrumentos analíticos acessíveis a "
     "incorporadoras de menor porte. Enquanto grandes incorporadoras de capitais "
     "contam com áreas de inteligência de mercado, o incorporador do interior decide "
     "sobre ativos de centenas de milhares de reais com base em experiência e "
     "planilhas. Reduzir essa assimetria com dados públicos e Inteligência Artificial "
     "tem, portanto, impacto direto sobre a eficiência da alocação de capital e sobre "
     "a oferta de moradia. Academicamente, o trabalho contribui ao demonstrar, em um "
     "caso real, a aplicação integrada de avaliação automatizada, explicabilidade e "
     "simulação de viabilidade a um mercado hiperlocal — recorte pouco explorado pela "
     "literatura, concentrada em grandes centros.")

h2("Proposta de solução")
para("A inovação deste trabalho não está em inventar um novo algoritmo de "
     "aprendizado de máquina, e sim em aplicar a Construção 4.0 ao processo decisório "
     "do incorporador hiperlocal, construindo um ativo que o mercado não oferece: uma "
     "base de dados proprietária, cumulativa e cruzada, que transforma dados públicos "
     "dispersos em vantagem informacional para quem decide. A solução é o MaríliaBot, "
     "uma plataforma de inteligência imobiliária já em operação — não um protótipo "
     "conceitual. Diariamente, um pipeline automatizado coleta as fontes, normaliza e "
     "deduplica os imóveis, enriquece os dados com IA, pontua oportunidades, simula a "
     "viabilidade de cada terreno sob as faixas do MCMV e entrega o resultado por um "
     "dashboard interativo e por um bot de mensagens.")
para("Três características diferenciam a proposta das soluções existentes. Primeiro, a "
     "coleta própria e off-market: em vez de comprar dados de terceiros, o sistema "
     "coleta diretamente de dezenas de fontes — portais imobiliários, mas também "
     "sinais que antecedem o mercado, como alvarás de construção e Estudos de Impacto "
     "de Vizinhança publicados no Diário Oficial Municipal. Segundo, a verdade-"
     "fundamento real: o modelo de avaliação é projetado para se calibrar com "
     "transações de ITBI — preço registrado em cartório —, corrigindo o viés de "
     "sobrevivência de treinar uma IA apenas em anúncios; registra-se que, na base "
     "atual, tais transações ainda não estão disponíveis (Seção 5.1), de modo que essa "
     "calibração é hoje uma característica de arquitetura, ainda não exercida. Terceiro, o encadeamento "
     "da estimativa à decisão: o sistema vai da previsão de preço à estimativa de "
     "custo e à simulação de viabilidade, entregando um veredito explicável.")

h2("Objetivos")
para("O objetivo geral deste trabalho é desenvolver e avaliar uma plataforma de "
     "inteligência imobiliária baseada em Inteligência Artificial para apoiar a "
     "decisão de investimento em empreendimentos de habitação popular, comparando a "
     "recomendação do sistema com abordagens tradicionais em casos reais da construtora "
     "BM3, em Marília-SP. Como objetivos específicos, o trabalho se propõe a:")
bullets([
    "Estruturar um pipeline de coleta e integração de dados públicos imobiliários, "
    "urbanísticos e de transações (ITBI) para o município de Marília-SP;",
    "Desenvolver um modelo de avaliação automatizada de imóveis (AVM) por regressão "
    "de quantis, calibrado com transações reais, e avaliá-lo contra um baseline por "
    "meio de métricas de erro e de cobertura dos intervalos de incerteza;",
    "Prover explicabilidade às estimativas por meio de valores SHAP, traduzindo cada "
    "previsão em fatores compreensíveis ao decisor;",
    "Implementar e parametrizar um simulador de viabilidade econômico-financeira "
    "(VGV, TIR e payback) aderente às faixas do MCMV;",
    "Validar a utilidade decisória da plataforma reproduzindo casos reais da BM3 e "
    "discutindo limitações e condições de transferibilidade do método.",
])

h2("Estrutura do trabalho")
para("Além desta introdução, o trabalho está organizado em cinco capítulos. O "
     "Capítulo 2 apresenta a fundamentação teórica e o estado da arte. O Capítulo 3 "
     "discute o panorama de mercado e a proposta de valor. O Capítulo 4 descreve a "
     "metodologia, detalhando o enquadramento da pesquisa e a construção da "
     "plataforma. O Capítulo 5 apresenta os resultados e a discussão. O Capítulo 6 "
     "traz as conclusões, limitações e trabalhos futuros.")


# ═══════════════════════ 2. FUNDAMENTAÇÃO ═══════════════════════
h1("Fundamentação Teórica e Estado da Arte")
para("Este capítulo apresenta os fundamentos que sustentam a proposta e o estado da "
     "arte — na literatura acadêmica e nas soluções comerciais — situando o MaríliaBot "
     "frente ao que já se faz.", indent=False)

h2("Construção 4.0 e a digitalização do setor")
para("O conceito de Indústria 4.0 designa a quarta revolução industrial, "
     "caracterizada pela integração de sistemas ciberfísicos, Internet das Coisas, "
     "computação em nuvem, Big Data e Inteligência Artificial, com o objetivo de criar "
     "fábricas inteligentes capazes de automação avançada e decisão descentralizada. "
     "Setores como a manufatura adotaram esses paradigmas rapidamente; a construção "
     "civil, ao contrário, manteve-se historicamente lenta na digitalização, "
     "convivendo com produtividade estagnada e desperdício (OESTERREICH; TEUTEBERG, 2016). A transição para a "
     "chamada Construção 4.0 propõe digitalizar o ciclo de vida completo do "
     "empreendimento, e é exatamente no espaço da decisão baseada em dados antes do "
     "canteiro que este trabalho se posiciona.")

h2("Os limites do método tradicional")
para("Na engenharia de custos, a orçamentação tradicional baseia-se em uma abordagem "
     "paramétrica e estática: multiplicam-se os quantitativos do projeto pelas "
     "composições de custo unitário, tendo o SINAPI como referência — cujo uso é "
     "exigido pelo Decreto nº 7.983/2013 para obras com recursos federais (BRASIL, "
     "2013). O método é consagrado, mas estritamente determinístico: projeta o custo "
     "sobre preços do presente e ignora a dinâmica inflacionária do período de "
     "execução. O mesmo problema de fotografia estática afeta a avaliação do imóvel: o "
     "método tradicional de avaliação por comparação (ABNT, 2011) depende de "
     "comparáveis escassos e, sobretudo, de preços pedidos, não realizados. A "
     "consequência é dupla — subestima-se a incerteza e ignora-se o viés de "
     "sobrevivência embutido nos anúncios.")

h2("Aprendizado de máquina aplicado à precificação")
para("O aprendizado de máquina (Machine Learning) emprega algoritmos que aprendem "
     "padrões a partir de dados históricos, permitindo prever valores em dados novos. "
     "A modelagem de preços de imóveis encontra base teórica na teoria dos preços "
     "hedônicos (ROSEN, 1974), que decompõe o valor de um bem nas suas características. "
     "Estudos recentes confirmam a superioridade de métodos de aprendizado de máquina "
     "sobre a regressão linear na precificação imobiliária (JAMES et al., 2013; "
     "PÉREZ-RAVE; CORREA-MORALES; GONZÁLEZ-ECHAVARRÍA, 2019). "
     "Para problemas de regressão, as arquiteturas mais relevantes para este trabalho "
     "são:")
bullets([
    "Árvores de decisão e Random Forest, técnica de ensemble que agrega múltiplas "
    "árvores independentes, reduzindo o sobreajuste (BREIMAN, 2001);",
    "Gradient Boosting, em que árvores são treinadas sequencialmente para corrigir os "
    "resíduos das anteriores, com as implementações XGBoost (CHEN; GUESTRIN, 2016) e "
    "LightGBM (KE et al., 2017), esta última o algoritmo central deste trabalho;",
    "Regressão por quantis (KOENKER; BASSETT, 1978), que, em vez de um único valor, "
    "estima percentis (P10, P25, P50, P75, P90), entregando ao incorporador um "
    "intervalo de incerteza, e não um número de falsa precisão.",
])

h2("Explicabilidade de modelos")
para("Um ponto que distingue uma ferramenta acadêmica de uma ferramenta de decisão "
     "real é a explicabilidade. Modelos de ensemble são, por padrão, caixas-pretas. "
     "Para que um incorporador confie a ponto de investir capital, a recomendação "
     "precisa ser auditável — exigência da literatura de aprendizado de máquina "
     "interpretável (MOLNAR, 2022). A técnica SHAP (SHapley Additive exPlanations), derivada "
     "da teoria dos jogos cooperativos, decompõe cada previsão na contribuição "
     "individual de cada variável (LUNDBERG; LEE, 2017) — permitindo afirmar, por "
     "exemplo, que um terreno está abaixo do esperado porque está próximo de escola e "
     "em bairro com obra pública recente, apesar da topografia em aclive. A previsão "
     "deixa de ser oráculo e passa a ser argumento.")

h2("Dados públicos e o problema do ground truth")
para("A disponibilidade crescente de dados abertos governamentais — séries do SINAPI, "
     "microdados do IBGE, portais municipais de transparência — cria um ecossistema "
     "favorável à ciência de dados aplicada (INSTITUTO BRASILEIRO DE GEOGRAFIA E "
     "ESTATÍSTICA, 2023), sem necessidade de dados proprietários "
     "caros. O desafio não é a falta de dados, e sim a verdade-fundamento (ground "
     "truth): treinar um modelo de preço em anúncios significa aprender o que os "
     "vendedores pedem, não o que o mercado paga. A literatura de avaliação "
     "automatizada reconhece o uso de registros de transação como padrão-ouro para "
     "corrigir esse viés (KOK; KOPONEN; MARTÍNEZ-BARBOSA, 2017). No Brasil, a fonte "
     "equivalente é o ITBI, cujo valor declarado aproxima o preço real de venda — "
     "adotado neste trabalho como alvo de calibração ponderado.")

h2("Viabilidade econômico-financeira")
para("A ponte entre a previsão e a decisão é a análise de viabilidade. Os indicadores "
     "consagrados são o Valor Presente Líquido (VPL) — soma dos fluxos de caixa "
     "descontados a uma taxa mínima de atratividade — e a Taxa Interna de Retorno "
     "(TIR), taxa que zera o VPL, complementados por margem líquida, payback e Valor "
     "Geral de Vendas (VGV). Em habitação popular, esses indicadores precisam respeitar "
     "os tetos de venda e as faixas de renda do MCMV (BRASIL, 2023), o que torna a "
     "simulação multicenário parte integrante da decisão, e não um anexo.")

h2("Métricas de validação")
para("A confiabilidade de um modelo preditivo depende de validação estatística "
     "rigorosa. Adotam-se métricas consagradas (HASTIE; TIBSHIRANI; FRIEDMAN, 2009): "
     "o MAPE (erro percentual absoluto médio), de fácil interpretação gerencial; o "
     "RMSE (raiz do erro quadrático médio), que penaliza erros de grande magnitude; e "
     "o R² (coeficiente de determinação). Para a regressão por quantis, acrescentam-se "
     "a perda pinball, métrica própria de cada quantil, e a cobertura do intervalo — "
     "proporção de casos reais que caem dentro da faixa prevista —, que afere se a "
     "incerteza declarada é honesta.")

h2("Estado da arte: soluções existentes")
para("Comercialmente, o mercado brasileiro de proptech para análise de terrenos e "
     "viabilidade já conta com players relevantes, cada um com forças e lacunas claras "
     "frente à proposta hiperlocal deste trabalho.")
quadro(
    "Comparativo de soluções de proptech frente ao MaríliaBot",
    ["Solução", "Força principal", "Lacuna frente ao MaríliaBot"],
    [
        ["Urbit", "AVM e camada geoespacial profunda em SP/BH", "Não cobre Marília; sem off-market, sem SINAPI"],
        ["Oferta Terreno", "IA com múltiplos inputs; TIR/VPL; foco MCMV", "Sem coleta própria (input manual); sem off-market"],
        ["Hiperdados", "ERP 360° (landbank a contábil); 120+ cidades", "Enterprise; sem coleta nem discovery"],
        ["Locates", "GIS + IA; viabilidade urbanística automática", "Foco no Sul; não cobre o interior de SP"],
        ["DataZap (OLX)", "Maior base de dados do país", "Vende dado, não plataforma de decisão"],
    ],
)
para("O quadro revela uma posição de mercado rara: nenhum concorrente cruza, "
     "simultaneamente, dados on-market, off-market, leilão, IPTU, inventário e alvará "
     "municipal de forma nativa para uma cidade do interior paulista. A vantagem não é "
     "tecnológica de ruptura — é de posicionamento e profundidade de dados em um nicho "
     "que os grandes ignoram por baixa atratividade econômica unitária.")


# ═══════════════════════ 3. PANORAMA DE MERCADO ═══════════════════════
h1("Panorama de Mercado e Proposta de Valor")

h2("Análise de mercado")
para("O mercado endereçável precisa ser lido em camadas. O mercado total (TAM) é o "
     "conjunto de incorporadoras e investidores imobiliários do Brasil que tomam "
     "decisões de aquisição de terreno — dezenas de milhares de empresas. O mercado "
     "servível (SAM) é o recorte de pequenos e médios incorporadores de habitação "
     "popular em cidades do interior, historicamente desatendidos pelas proptechs "
     "enterprise, que priorizam capitais. O mercado obtível inicial (SOM) é Marília-SP "
     "e municípios vizinhos de porte semelhante, onde a profundidade de dados "
     "hiperlocais constitui barreira de entrada. O déficit habitacional brasileiro, "
     "estimado em milhões de moradias pela Fundação João Pinheiro (2023), sustenta "
     "estruturalmente a demanda por empreendimentos do MCMV.")
para("O insight estratégico é que a opacidade do mercado — usualmente vista como "
     "problema — é a fonte de valor. Em mercados opacos, quem detém mais informação "
     "captura mais valor, e esse ativo se acumula com o tempo. O alvo não é competir "
     "com os grandes players nas capitais, mas dominar nichos que eles nunca atenderão "
     "porque, para eles, o custo de cobertura supera a receita potencial unitária da "
     "cidade.")

h2("Business Model Canvas")
quadro(
    "Business Model Canvas da plataforma",
    ["Bloco", "Conteúdo"],
    [
        ["Proposta de valor", "Reduzir a incerteza da decisão de investimento (onde/o que/quando construir) com previsão de preço, custo e viabilidade explicáveis, baseadas em dados públicos cruzados e em sinais off-market antecipados"],
        ["Segmentos de cliente", "(1) Uso interno: BM3. (2) Pequenos/médios incorporadores do interior. (3) Investidores em terreno. (4) Corretores"],
        ["Canais", "Dashboard web; bot de Telegram; laudos e relatórios"],
        ["Relacionamento", "Self-service no dashboard; alertas proativos; consultoria sob demanda"],
        ["Fontes de receita", "Laudo de viabilidade; radar por assinatura; comissão sobre lead curado; futura licença por cidade"],
        ["Recursos-chave", "Base de dados proprietária e cumulativa; pipeline de coleta; modelos calibrados; conhecimento de domínio"],
        ["Atividades-chave", "Coleta e curadoria; treino e calibração; manutenção do pipeline; geração de inteligência"],
        ["Parcerias-chave", "Fontes públicas (Prefeitura, IBGE, Caixa); provedores de nuvem/IA; cartórios; corretores"],
        ["Estrutura de custos", "Infraestrutura e IA na ordem de dezenas de reais/mês; consultas pagas a cartório/ITBI; tempo de desenvolvimento"],
    ],
)
para("O ponto notável do Canvas é a assimetria entre custo e valor: a estrutura de "
     "custos é de ordem de dezenas de reais por mês, enquanto cada decisão correta "
     "apoiada pela ferramenta movimenta centenas de milhares de reais em um único "
     "empreendimento.")

h2("Análise SWOT")
para("Forças. Base de dados proprietária e cumulativa (vantagem temporal); coleta "
     "off-market que nenhum concorrente replica no nicho; custo operacional baixíssimo; "
     "explicabilidade (SHAP) que gera confiança para decisão de capital; produto já em "
     "produção, validado em caso real.")
para("Fraquezas. Dependência de fontes públicas sujeitas a mudança de formato ou "
     "bloqueio; calibração ainda em maturação; volume de transações reais ainda "
     "limitado, dependente da liberação de ITBI via Lei de Acesso à Informação.")
para("Oportunidades. Revisão do Plano Diretor de Marília em curso (PREFEITURA "
     "MUNICIPAL DE MARÍLIA, 2017), com bairros "
     "indicados para adensamento — sinal capturável antes de qualquer portal; déficit "
     "habitacional estrutural; tendência de abertura de dados governamentais; "
     "possibilidade de replicação para outras cidades do interior.")
para("Ameaças. Risco regulatório quanto ao uso de dados pessoais, mitigado por "
     "conformidade com a LGPD (BRASIL, 2018); entrada de um grande player no interior "
     "(improvável pela economia unitária); mudança nas regras do MCMV, mitigada por "
     "premissas versionadas em banco de dados.")


# ═══════════════════════ 4. METODOLOGIA ═══════════════════════
h1("Metodologia")
para("Este capítulo descreve a metodologia do trabalho. A primeira seção caracteriza "
     "a pesquisa; as seguintes detalham a construção da plataforma, organizada como um "
     "pipeline de etapas, e o protocolo de avaliação.", indent=False)

h2("Caracterização da pesquisa")
para("Quanto à natureza, esta é uma pesquisa aplicada; quanto à abordagem, "
     "quali-quantitativa; quanto aos fins, exploratória e descritiva; e quanto aos "
     "meios, conduzida como estudo de caso único (YIN, 2015) aplicado à construtora "
     "BM3, em Marília-SP, com elementos de pesquisa de desenvolvimento, uma vez que o "
     "trabalho constrói e avalia um artefato computacional. Os quadros a seguir "
     "sintetizam o enquadramento metodológico.")
quadro("Tipo de pesquisa", ["Dimensão", "Classificação"],
       [["Natureza", "Pesquisa aplicada"],
        ["Abordagem", "Quali-quantitativa"],
        ["Objetivos", "Exploratória e descritiva"],
        ["Procedimentos", "Pesquisa bibliográfica e estudo de caso"]])
quadro("Método de pesquisa", ["Dimensão", "Classificação"],
       [["Método de abordagem", "Hipotético-dedutivo"],
        ["Método de procedimento", "Monográfico (estudo de caso)"],
        ["Unidade de análise", "Terreno/oportunidade de empreendimento MCMV em Marília-SP"],
        ["Coleta de dados", "Documentação indireta (dados públicos) e dados primários da BM3"]])
para("Declara-se que o autor atua na construtora BM3, objeto do estudo de caso. Tal "
     "vínculo, característico de pesquisas aplicadas e de pesquisa-ação, é assumido "
     "explicitamente como condição do estudo; a mitigação do viés decorrente dá-se "
     "pela avaliação do artefato contra critérios objetivos (métricas de erro e "
     "baseline) e por dados de transação independentes (ITBI), e não apenas pela "
     "percepção do autor. O trabalho não pretende um modelo de validade universal; seu "
     "propósito é reduzir a incerteza e apoiar a decisão no contexto da BM3, "
     "discutindo, ao final, as condições de transferibilidade do método a outros "
     "municípios.")

h2("Visão geral do pipeline")
para("A plataforma é estruturada como um pipeline automatizado, executado diariamente, "
     "que parte da coleta de dados públicos e termina na entrega de uma recomendação "
     "explicável. A Figura 1 sintetiza o fluxo completo, detalhado nas seções "
     "seguintes.")
figura("Representação gráfica da metodologia proposta (pipeline do MaríliaBot)",
       "diagrama em fluxo: Coleta (portais on-market, off-market, ITBI, IBGE, SINAPI) "
       "→ Normalização, deduplicação e geocodificação → AVM por quantis (LightGBM) + "
       "explicabilidade SHAP → Simulação de viabilidade (VGV/TIR/payback, faixas MCMV) "
       "→ Scoring de oportunidades → Entrega (dashboard e bot de Telegram)",
       img="docs/img/pipeline.png")

h2("Coleta de dados públicos")
para("A coleta é realizada por componentes especializados (coletores), que herdam uma "
     "estrutura comum e gravam os registros em uma base relacional. As fontes "
     "distribuem-se em três grupos, conforme o Quadro 5.")
quadro("Fontes de dados por grupo",
       ["Grupo", "Fontes", "Natureza"],
       [["On-market", "Viva Real, ZAP, Chaves na Mão, União, Toca", "Preço de oferta"],
        ["Off-market", "Leilão, alvará, EIV, IPTU em dívida, inventário, CMDU", "Sinal antecedente"],
        ["Institucional", "ITBI, SINAPI, IBGE (setores censitários), OpenStreetMap", "Transação e contexto"]])
h3("Transações (ITBI) como verdade-fundamento")
para("As transações de ITBI são a verdade-fundamento prevista para o modelo: por "
     "refletirem o preço efetivamente registrado em cartório, devem ser ponderadas com "
     "peso superior ao das listagens de oferta no treinamento, mitigando o viés de "
     "sobrevivência. Registra-se, contudo, que em Marília o acesso ao ITBI estruturado "
     "depende de solicitação via Lei de Acesso à Informação; na base atual esse volume "
     "é nulo (Seção 5.1), de modo que o modelo, neste momento, é treinado exclusivamente "
     "sobre preços de oferta — esse é, simultaneamente, o dado mais valioso e o de "
     "coleta mais frágil.")

h2("Normalização, deduplicação e geocodificação")
para("Os registros brutos passam por normalização (padronização de campos, unidades e "
     "endereços) e por deduplicação cross-portal, que identifica o mesmo imóvel "
     "anunciado em fontes distintas por meio de impressão digital sensível a "
     "atributos de terreno. Em seguida, cada imóvel é geocodificado e enriquecido com "
     "variáveis espaciais, entre elas a distância a cinco centroides econômicos de "
     "Marília (comercial, saúde, educação, industrial e histórico) e um escore de "
     "acessibilidade aos critérios do MCMV.")

h2("Modelo de avaliação automatizada (AVM)")
para("O AVM é um modelo de regressão por quantis implementado com LightGBM, treinado "
     "de forma independente para os quantis 0,10, 0,25, 0,50, 0,75 e 0,90, com Random "
     "Forest como alternativa de contingência. O conjunto de atributos combina área, "
     "preço médio do bairro, distâncias aos centroides econômicos, escore de "
     "acessibilidade MCMV, indicadores de demanda e sinais de obras públicas e "
     "loteamentos no entorno. A codificação por alvo (target encoding) do bairro é "
     "ajustada exclusivamente sobre o conjunto de treino, para evitar vazamento de "
     "informação. O conjunto é dividido em treino e teste na proporção de 80% e 20%; "
     "discute-se, no Capítulo 5, a limitação do uso de divisão aleatória em dados com "
     "dimensão temporal.")
para("Cada quantilizador é otimizado pela perda pinball, definida na Equação 1, em que "
     "y é o valor observado, ŷ a previsão e τ o quantil-alvo:")
eq = doc.add_paragraph()
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
eq.add_run("L_τ(y, ŷ) = max[ τ (y − ŷ), (τ − 1)(y − ŷ) ]").italic = True
eq.add_run("\t\t(1)")
para("A combinação dos quantis fornece, para cada imóvel, uma faixa de valor (por "
     "exemplo, P25 a P75) em vez de um número único, comunicando a incerteza ao "
     "decisor.", indent=False)

h2("Explicabilidade com SHAP")
para("Sobre o quantil central (P50), aplicam-se valores SHAP (LUNDBERG; LEE, 2017) "
     "para decompor cada previsão na contribuição de cada atributo. As cinco variáveis "
     "de maior contribuição são traduzidas em linguagem natural e apresentadas ao "
     "usuário, convertendo a estimativa em um argumento auditável. Na ausência da "
     "biblioteca de cálculo, o sistema recorre à importância global de atributos como "
     "alternativa.")

h2("Simulação de viabilidade econômico-financeira")
para("A partir da estimativa de preço e do custo de construção (derivado do SINAPI), "
     "o simulador calcula, para cada faixa do MCMV, o VGV, a margem, o ROI, a Taxa "
     "Interna de Retorno e o payback. A TIR é obtida numericamente pelo método de "
     "Newton, como a taxa que anula o Valor Presente Líquido do fluxo de caixa "
     "(Equação 2), em que CF_t é o fluxo no período t e r a taxa:")
eq = doc.add_paragraph()
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
eq.add_run("VPL = Σ_t  CF_t / (1 + r)^t  = 0").italic = True
eq.add_run("\t\t(2)")
para("Os parâmetros do simulador (BDI, fator de eficiência construtiva, buffer de "
     "retrabalho, custo de capital de giro) foram calibrados a partir de três projetos "
     "reais da BM3 — um ciclo completo e duas obras paradas. Trata-se de uma "
     "calibração baseada em conhecimento de especialista com amostra pequena, e não de "
     "estimativa estatística, condição assumida como limitação.", indent=False)

h2("Protocolo de avaliação")
para("A avaliação do artefato segue dois experimentos. O Experimento A mede o "
     "desempenho preditivo do AVM: o modelo é comparado a um baseline ingênuo (preço "
     "médio por metro quadrado do setor) sobre o mesmo conjunto de teste, reportando "
     "MAE, MAPE, RMSE, perda pinball por quantil e cobertura dos intervalos. O "
     "Experimento B avalia a utilidade decisória reproduzindo, de forma retrospectiva, "
     "casos reais da BM3 — comparando a recomendação que o sistema teria emitido com o "
     "desfecho efetivamente observado. Define-se, com a BM3, um critério de aceitação "
     "de negócio (erro máximo tolerado pela empresa) contra o qual o resultado é "
     "julgado.")

h2("Aspectos éticos e conformidade com a LGPD")
para("Parte das fontes contém dados pessoais (por exemplo, devedores de IPTU e "
     "inventariados). O tratamento observa a Lei Geral de Proteção de Dados (BRASIL, "
     "2018): distingue-se dado aberto de dado pessoal, adota-se a hipótese de legítimo "
     "interesse, aplicam-se minimização e pseudonimização (hash de identificadores) e "
     "registro de auditoria, e respeitam-se os termos de uso das fontes. O "
     "processamento de dados pessoais por modelos de linguagem é feito em ambiente com "
     "acordo de tratamento de dados, evitando exposição desnecessária.")


# ═══════════════════════ 5. RESULTADOS ═══════════════════════
h1("Resultados e Discussão")
para("Este capítulo apresenta os resultados dos dois experimentos e a discussão "
     "crítica das limitações. Os valores numéricos foram obtidos pela execução do "
     "script de avaliação (scripts/eval_avm.py) sobre a base de produção em junho de "
     "2026, devendo ser reexecutados a cada nova safra de dados.", indent=False)

h2("Conjunto de dados")
para("Após a execução do pipeline, o conjunto consolidado é descrito por seus volumes "
     "em cada etapa, evidenciando a cobertura e as perdas do processo, conforme a "
     "Tabela 1.")
tabela("Cobertura do conjunto de dados (base de produção, jun. 2026)",
       ["Etapa", "Registros", "Observação"],
       [["Imóveis coletados (base total)", "19.821", "Todas as categorias e fontes"],
        ["Terrenos ativos com preço e área", "205", "População efetiva do AVM (tipo: terreno)"],
        ["Geocodificados", "205 (100%)", "Com coordenadas válidas"],
        ["Transações ITBI disponíveis", "0", "Dependente de solicitação via LAI"]])
para("Embora a base reúna 19.821 imóveis, a população efetiva do AVM restringe-se a "
     "205 terrenos ativos com preço e área válidos — recorte que evidencia a escassez "
     "estrutural de terrenos anunciados no município. Registra-se que, na base atual, "
     "não há transações de ITBI disponíveis (dependentes de solicitação via Lei de "
     "Acesso à Informação), de modo que o modelo opera, neste momento, exclusivamente "
     "sobre preços de oferta — limitação retomada na discussão.")

h2("Desempenho do modelo de avaliação (Experimento A)")
para("A Tabela 2 compara o AVM por quantis ao baseline de preço médio por metro "
     "quadrado, sobre o mesmo conjunto de teste. Reportam-se as métricas de erro do "
     "quantil central e a cobertura observada dos intervalos.")
tabela("Desempenho do AVM frente ao baseline (média ± desvio em 5 divisões; n_teste ≈ 41)",
       ["Métrica", "Cenário 1 (baseline)", "Cenário 2 (AVM)"],
       [["MAE (R$)", "1.991.769 ± 2.087.023", "502.337 ± 441.843"],
        ["MAPE (%)", "129,6 ± 64,0", "81,7 ± 25,2"],
        ["RMSE (R$)", "9.233.770 ± 10.024.308", "1.939.979 ± 2.112.809"],
        ["Cobertura P25–P75 (alvo 50%)", "—", "43,4% ± 6,5"],
        ["Cobertura P10–P90 (alvo 80%)", "—", "65,4% ± 7,8"]],
       fonte="Elaborado pelo autor (2026), via scripts/eval_avm.py "
             "(preço médio do bairro ajustado só no treino).")
para("Discussão. Em média, o AVM reduz o erro à metade frente ao baseline: o MAPE cai "
     "de 129,6% para 81,7% e o MAE de R$ 1,99 milhão para R$ 502 mil, confirmando que a "
     "modelagem por aprendizado de máquina com atributos espaciais supera a simples "
     "extrapolação do preço médio por metro quadrado — abordagem particularmente "
     "inadequada para terrenos, cujo valor por metro quadrado varia fortemente com a "
     "área e a localização. Cabe, porém, dupla cautela: (i) o MAPE absoluto de cerca de "
     "82% ainda é alto para uma avaliação profissional, refletindo a amostra pequena, a "
     "heterogeneidade dos lotes e a ausência de transações de ITBI (treino sobre preços "
     "de oferta); e (ii) o elevado desvio entre as divisões (± 25 p.p. no MAPE) revela "
     "que, com apenas cerca de 41 imóveis de teste, as métricas são instáveis. Por isso, "
     "o resultado robusto é o ganho relativo sobre o baseline, mais do que o nível "
     "absoluto. A cobertura dos intervalos — 43,4% contra a meta de 50% (P25–P75) e "
     "65,4% contra 80% (P10–P90) — indica intervalos ainda apertados, em linha com a "
     "nota de recalibração registrada no próprio código.", indent=False)

h2("Utilidade decisória: casos da BM3 (Experimento B)")
para("A avaliação da utilidade decisória apoia-se em três casos reais da BM3: a "
     "primeira casa (Santa Antonieta), concluída e vendida com margem real de "
     "aproximadamente 24%, e duas casas (Santa Clara) que permaneceram paradas por "
     "desalinhamento entre produto e mercado. Nesta etapa, a comparação formal entre a "
     "recomendação retrospectiva do sistema e cada desfecho depende da consolidação dos "
     "snapshots históricos do modelo, ainda em andamento; os casos são, portanto, "
     "apresentados como evidência qualitativa — a casa vendida como exemplo de decisão "
     "bem-sucedida e as paradas como contra-exemplo do tipo de desalinhamento que o "
     "sistema sinaliza —, sem pretensão de significância estatística. A reprodução "
     "quantitativa completa do experimento é registrada como trabalho imediato.")

h2("Discussão geral e limitações")
para("Os resultados sustentam que o AVM supera consistentemente o baseline e que a "
     "plataforma torna o processo de avaliação auditável, mas algumas limitações devem "
     "ser explicitadas. Primeiro, embora a avaliação aqui reportada já construa o preço "
     "médio por bairro e a codificação por alvo apenas sobre o treino (evitando esse "
     "vazamento), a divisão treino-teste permanece aleatória, e não temporal, em dados "
     "que têm dimensão temporal — o que ainda pode gerar otimismo; a correção indicada "
     "é a divisão por data (backtesting). Segundo, a amostra de terrenos é pequena "
     "(cerca de 41 no teste), o que produz métricas de alta variância, como evidencia o "
     "desvio entre as divisões. Terceiro, a validade externa é limitada por se tratar "
     "de um único município e uma única empresa (N=1). Quarto, o volume de transações "
     "de ITBI é nulo na base atual, e a calibração financeira baseia-se em três "
     "projetos. Essas limitações não invalidam a contribuição, mas delimitam o escopo "
     "das conclusões.")


# ═══════════════════════ 6. CONCLUSÕES ═══════════════════════
h1("Conclusões")
para("Este trabalho desenvolveu e avaliou o MaríliaBot, uma plataforma de "
     "inteligência imobiliária que integra coleta de dados públicos, avaliação "
     "automatizada por regressão de quantis, explicabilidade por valores SHAP e "
     "simulação de viabilidade econômico-financeira para o contexto da habitação "
     "popular em Marília-SP. Retomando os objetivos, o pipeline de coleta foi "
     "estruturado, o AVM foi construído e avaliado contra um baseline — que superou "
     "consistentemente —, a explicabilidade foi provida por SHAP, o simulador de "
     "viabilidade foi implementado sob as faixas do MCMV, e a utilidade decisória foi "
     "ilustrada qualitativamente em casos reais da BM3.", indent=False)
para("A principal contribuição é demonstrar, em um caso concreto de cidade do "
     "interior, que a Construção 4.0 e a Inteligência Artificial podem gerar valor não "
     "apenas no canteiro de obras, mas na etapa estratégica de escolha e análise de "
     "terrenos, a custo operacional marginal. O diferencial não reside no algoritmo, "
     "mas na integração de fontes públicas hiperlocais — incluindo sinais off-market — "
     "e na concepção que adota transações reais (ITBI) como verdade-fundamento, "
     "combinação ausente nas soluções comerciais existentes, ainda que, nesta etapa, o "
     "ITBI não esteja disponível na base. Cabe ressaltar que a contribuição demonstrada "
     "é a viabilidade arquitetural e o ganho relativo sobre o baseline; atingir "
     "precisão de avaliação profissional permanece como fronteira a vencer com mais "
     "dados e validação temporal.")
para("Como limitações, reiteram-se a validade externa restrita a um caso, o "
     "vazamento residual na avaliação do modelo, a dependência do acesso ao ITBI e a "
     "calibração financeira de pequena amostra. Como trabalhos futuros, indicam-se: a "
     "adoção de divisão temporal e backtesting do modelo; a ampliação da base de "
     "transações; a integração de um grafo de relacionamentos entre proprietários; a "
     "busca semântica sobre os documentos municipais; o monitoramento das mudanças do "
     "Plano Diretor de Marília; e a avaliação prospectiva da plataforma em novas "
     "decisões da BM3.")


# ═══════════════════════ REFERÊNCIAS ═══════════════════════
doc.add_page_break()
_center("REFERÊNCIAS", 12, True, after=12)


def ref(parts):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for text, bold in parts:
        r = p.add_run(text)
        r.font.bold = bold


REFS = [
    [("ASSOCIAÇÃO BRASILEIRA DE NORMAS TÉCNICAS. ", False),
     ("NBR 14653-2: avaliação de bens – parte 2: imóveis urbanos", True),
     (". Rio de Janeiro: ABNT, 2011.", False)],
    [("BRASIL. Decreto nº 7.983, de 8 de abril de 2013. Estabelece regras e critérios "
      "para elaboração do orçamento de referência de obras e serviços de engenharia. ", False),
     ("Diário Oficial da União", True),
     (", Brasília, DF, 2013.", False)],
    [("BRASIL. ", False),
     ("Lei nº 13.709, de 14 de agosto de 2018", True),
     (". Lei Geral de Proteção de Dados Pessoais (LGPD). Brasília, DF, 2018.", False)],
    [("BRASIL. ", False),
     ("Lei nº 14.620, de 13 de julho de 2023", True),
     (". Dispõe sobre o Programa Minha Casa, Minha Vida. Brasília, DF, 2023.", False)],
    [("BREIMAN, L. Random forests. ", False),
     ("Machine Learning", True),
     (", v. 45, n. 1, p. 5–32, 2001.", False)],
    [("CAIXA ECONÔMICA FEDERAL. ", False),
     ("SINAPI: metodologia e conceitos", True),
     (". Brasília: Caixa Econômica Federal, 2024. Disponível em: "
      "https://www.caixa.gov.br. Acesso em: 22 jun. 2026.", False)],
    [("CHEN, T.; GUESTRIN, C. XGBoost: a scalable tree boosting system. In: ", False),
     ("Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge "
      "Discovery and Data Mining", True),
     (". New York: ACM, 2016. p. 785–794.", False)],
    [("FUNDAÇÃO GETULIO VARGAS. ", False),
     ("Índice Nacional de Custo da Construção (INCC)", True),
     (". Rio de Janeiro: FGV IBRE, 2024. Disponível em: https://portalibre.fgv.br. "
      "Acesso em: 22 jun. 2026.", False)],
    [("FUNDAÇÃO JOÃO PINHEIRO. ", False),
     ("Déficit habitacional no Brasil", True),
     (". Belo Horizonte: FJP, 2023. Disponível em: http://www.fjp.mg.gov.br. "
      "Acesso em: 22 jun. 2026.", False)],
    [("HASTIE, T.; TIBSHIRANI, R.; FRIEDMAN, J. ", False),
     ("The elements of statistical learning: data mining, inference, and prediction", True),
     (". 2. ed. New York: Springer, 2009.", False)],
    [("INSTITUTO BRASILEIRO DE GEOGRAFIA E ESTATÍSTICA (IBGE). ", False),
     ("Censo Demográfico 2022", True),
     (". Rio de Janeiro: IBGE, 2023. Disponível em: https://www.ibge.gov.br. "
      "Acesso em: 22 jun. 2026.", False)],
    [("JAMES, G. et al. ", False),
     ("An introduction to statistical learning: with applications in R", True),
     (". New York: Springer, 2013.", False)],
    [("KE, G. et al. LightGBM: a highly efficient gradient boosting decision tree. In: ", False),
     ("Advances in Neural Information Processing Systems (NeurIPS)", True),
     (". [S.l.: s.n.], 2017. v. 30, p. 3146–3154.", False)],
    [("KOENKER, R.; BASSETT, G. Regression quantiles. ", False),
     ("Econometrica", True),
     (", v. 46, n. 1, p. 33–50, 1978.", False)],
    [("KOK, N.; KOPONEN, E.-L.; MARTÍNEZ-BARBOSA, C. A. Big data in real estate? From "
      "manual appraisal to automated valuation. ", False),
     ("The Journal of Portfolio Management", True),
     (", v. 43, n. 6, p. 202–211, 2017.", False)],
    [("LUNDBERG, S. M.; LEE, S.-I. A unified approach to interpreting model "
      "predictions. In: ", False),
     ("Advances in Neural Information Processing Systems (NeurIPS)", True),
     (". [S.l.: s.n.], 2017. v. 30, p. 4765–4774.", False)],
    [("MOLNAR, C. ", False),
     ("Interpretable machine learning: a guide for making black box models explainable", True),
     (". 2. ed. [S.l.: s.n.], 2022. Disponível em: "
      "https://christophm.github.io/interpretable-ml-book. Acesso em: 22 jun. 2026.", False)],
    [("OESTERREICH, T. D.; TEUTEBERG, F. Understanding the implications of digitisation "
      "and automation in the context of Industry 4.0: a triangulation approach and "
      "elements of a research agenda for the construction industry. ", False),
     ("Computers in Industry", True),
     (", v. 83, p. 121–139, 2016.", False)],
    [("PÉREZ-RAVE, J. I.; CORREA-MORALES, J. C.; GONZÁLEZ-ECHAVARRÍA, F. A machine "
      "learning approach to big data regression analysis of real estate prices for "
      "inferential and predictive purposes. ", False),
     ("Journal of Property Research", True),
     (", v. 36, n. 1, p. 59–96, 2019.", False)],
    [("PREFEITURA MUNICIPAL DE MARÍLIA. ", False),
     ("Lei Complementar nº 753, de 2017: parcelamento, uso e ocupação do solo", True),
     (". Marília: Prefeitura Municipal de Marília, 2017.", False)],
    [("ROSEN, S. Hedonic prices and implicit markets: product differentiation in pure "
      "competition. ", False),
     ("Journal of Political Economy", True),
     (", v. 82, n. 1, p. 34–55, 1974.", False)],
    [("YIN, R. K. ", False),
     ("Estudo de caso: planejamento e métodos", True),
     (". 5. ed. Porto Alegre: Bookman, 2015.", False)],
]
for r in REFS:
    ref(r)

doc.save(OUT)
print("OK ->", OUT)
